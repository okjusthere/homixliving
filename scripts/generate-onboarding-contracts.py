#!/usr/bin/env python3
"""Author two legal-review masters and generate four entity-specific releases."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader, PdfWriter


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "contracts" / "source"
GENERATED_DIR = ROOT / "contracts" / "generated"
PDF_DIR = ROOT / "output" / "pdf"
ENTITIES_PATH = ROOT / "contracts" / "entities.yml"
MANIFEST_PATH = ROOT / "contracts" / "field-manifests.yml"
AGENT_MASTER = SOURCE_DIR / "Agent_Affiliation_Agreement.docx"
TEAM_LEADER_MASTER = SOURCE_DIR / "Team_Leader_Agreement.docx"
AGENT_BASELINE_PATH = SOURCE_DIR / "agent-affiliation-baseline.json"
LIBOR_APPLICATION_PATH = ROOT / "contracts" / "appendices" / "LIBOR_REALTOR_Application_Rev_10-25.pdf"
REALTY_FEE_DISCLOSURE_PATH = ROOT / "contracts" / "appendices" / "Realty_LIBOR_OneKey_Fee_Disclosures_v1.pdf"

INK, MUTED, LINE, PAPER, BRONZE, GREEN = "1D1C19", "6F6A61", "D8D1C5", "F4F1EA", "98623C", "536B3A"

AGENT_PLANS = {
    "solo": {
        "label": "Solo",
        "filename": "Solo",
        "economics": "85% Agent / 15% Company",
        "fee": "$288 / 12 months or $500 / 24 months",
        "cap_rule": "$12,000 annual Company Dollar cap",
        "terms": (
            "Solo: 85% Agent / 15% Company; $288 for 12 months or $500 for 24 months; "
            "$12,000 annual Company Dollar cap."
        ),
    },
    "solo_pro": {
        "label": "Solo Pro",
        "filename": "Solo_Pro",
        "economics": "100% Agent / 0% Company split",
        "fee": "$3,650 annual plan fee",
        "cap_rule": "Transaction fee schedule; no split cap",
        "terms": (
            "Solo Pro: 100% commission mode; $3,650 annual plan fee; transaction fee of $200 "
            "for a commission check between $10,000 and $30,000, $500 between $30,000 and "
            "$100,000, and $1,000 above $100,000."
        ),
    },
    "team_member": {
        "label": "Team Member",
        "filename": "Team_Member",
        "economics": "90% Agent Side / 10% Company",
        "fee": "$288 / 12 months or $500 / 24 months",
        "cap_rule": "$10,000 Company cap; Team terms separate",
        "terms": (
            "Team Member: 90% Agent Side / 10% Company; $288 for 12 months or $500 for "
            "24 months; $10,000 annual Company Dollar cap. Team Split and Team Cap are "
            "governed separately by the accepted Team Compensation Configuration."
        ),
    },
}


def load_json_yaml(path: Path) -> dict:
    # JSON is valid YAML 1.2 and avoids a runtime PyYAML dependency.
    return json.loads(path.read_text(encoding="utf-8"))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent=120) -> None:
    if sum(widths) != 9360:
        raise ValueError(f"Table widths must total 9360 DXA, got {sum(widths)}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
    if tbl_w.getparent() is None:
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
    if tbl_ind.getparent() is None:
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
            if tc_w.getparent() is None:
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(run, name="Times New Roman", size=9, bold=False, color=INK, italic=False) -> None:
    run.font.name = name
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_field(paragraph) -> None:
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    nodes = []
    for kind in ("begin", "separate", "end"):
        node = OxmlElement("w:fldChar")
        node.set(qn("w:fldCharType"), kind)
        nodes.append(node)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.extend([nodes[0], instr, nodes[1], nodes[2]])


def configure_document(doc: Document, running_label: str) -> None:
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin, section.right_margin = Inches(0.82), Inches(1)
    section.bottom_margin, section.left_margin = Inches(0.78), Inches(1)
    section.header_distance = section.footer_distance = Inches(0.36)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name, normal.font.size = "Times New Roman", Pt(9)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.paragraph_format.space_before, normal.paragraph_format.space_after = Pt(0), Pt(4)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, before, after, color in (
        ("Title", 22, 0, 6, INK), ("Subtitle", 10, 0, 14, MUTED),
        ("Heading 1", 12.5, 10, 5, INK), ("Heading 2", 10.5, 7, 3, BRONZE),
    ):
        style = styles[name]
        style.font.name, style.font.size = "Arial", Pt(size)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before, style.paragraph_format.space_after = Pt(before), Pt(after)
        style.paragraph_format.keep_with_next = True
    style = styles.add_style("Legal Small", WD_STYLE_TYPE.PARAGRAPH)
    style.font.name, style.font.size = "Arial", Pt(7.5)
    style.font.color.rgb = RGBColor.from_string(MUTED)
    style.paragraph_format.space_after, style.paragraph_format.line_spacing = Pt(3), 1.05
    list_style = styles["List Bullet"]
    list_style.font.name, list_style.font.size = "Times New Roman", Pt(9)
    list_style.paragraph_format.left_indent = Inches(0.38)
    list_style.paragraph_format.first_line_indent = Inches(-0.19)
    list_style.paragraph_format.space_after, list_style.paragraph_format.line_spacing = Pt(3), 1.08
    header = section.header.paragraphs[0]
    run = header.add_run(f"HOMIX  |  {running_label}")
    set_run_font(run, "Arial", 7.5, True, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_field(footer)
    for run in footer.runs:
        set_run_font(run, "Arial", 7.5, color=MUTED)


def add_para(doc: Document, text: str, *, bold=False, italic=False, small=False, align=None):
    paragraph = doc.add_paragraph(style="Legal Small" if small else "Normal")
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, "Arial" if small else "Times New Roman", 7.5 if small else 9,
                 bold, MUTED if small else INK, italic)
    return paragraph


def add_heading(doc: Document, text: str, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        set_run_font(paragraph.add_run(item))


def add_title_block(doc: Document, title: str, subtitle: str) -> None:
    paragraph = doc.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(paragraph.add_run(title), "Arial", 22, True)
    paragraph = doc.add_paragraph(style="Subtitle")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(paragraph.add_run(subtitle), "Arial", 10, color=MUTED)


def add_key_values(doc: Document, items: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=4)
    set_table_geometry(table, [1600, 3080, 1600, 3080])
    for index in range(0, len(items), 2):
        row = table.add_row()
        pair = items[index:index + 2]
        for pair_index in range(2):
            label_cell, value_cell = row.cells[pair_index * 2], row.cells[pair_index * 2 + 1]
            label, value = pair[pair_index] if pair_index < len(pair) else ("", "")
            set_cell_shading(label_cell, PAPER)
            label_cell.vertical_alignment = value_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_run_font(label_cell.paragraphs[0].add_run(label), "Arial", 7.5, True, MUTED)
            set_run_font(value_cell.paragraphs[0].add_run(value), "Arial", 8)


def add_plan_table(doc: Document) -> None:
    rows = [
        ("{{PLAN_LABEL}}", "{{PLAN_ECONOMICS}}", "{{PLAN_FEE}}", "{{PLAN_CAP_RULE}}"),
    ]
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [1500, 2420, 2560, 2880])
    for cell, text in zip(table.rows[0].cells, ["Plan", "Company economics", "Affiliation fee", "Cap / transaction rule"]):
        set_cell_shading(cell, PAPER)
        set_run_font(cell.paragraphs[0].add_run(text), "Arial", 7.5, True)
    for row_data in rows:
        row = table.add_row()
        for index, value in enumerate(row_data):
            set_run_font(row.cells[index].paragraphs[0].add_run(value), "Arial", 7.5,
                         index == 0, GREEN if index == 0 else INK)
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=45, start=120, bottom=45, end=120)


def add_signature_grid(doc: Document, left_label: str, right_label: str) -> None:
    table = doc.add_table(rows=4, cols=2)
    set_table_geometry(table, [4680, 4680])
    values = [
        (left_label, right_label),
        ("Signature: __________________________________", "Signature: __________________________________"),
        ("Printed name: ______________________________", "{{BROKER_NAME}}, {{BROKER_TITLE}}"),
        ("Date: ______________________________________", "Date: ______________________________________"),
    ]
    for row_index, row_values in enumerate(values):
        for col_index, value in enumerate(row_values):
            cell = table.rows[row_index].cells[col_index]
            if row_index == 0:
                set_cell_shading(cell, PAPER)
            set_run_font(cell.paragraphs[0].add_run(value), "Arial", 7.5, row_index == 0)
            set_cell_margins(cell, top=55, start=120, bottom=55, end=120)


def add_acknowledgement(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_run_font(paragraph.add_run("[  ] "), "Arial", 10, True, GREEN)
    set_run_font(paragraph.add_run(text), "Arial", 8.5, True)


def new_page(doc: Document, title: str, kicker: str) -> None:
    doc.add_page_break()
    paragraph = doc.add_paragraph(kicker.upper())
    paragraph.paragraph_format.space_after = Pt(2)
    set_run_font(paragraph.runs[0], "Arial", 7.5, True, BRONZE)
    add_heading(doc, title)


def add_legal_section(doc: Document, title: str, paragraphs: list[str]) -> None:
    add_heading(doc, title, 2)
    for paragraph in paragraphs:
        add_para(doc, paragraph)


def source_legal_text(text: str) -> str:
    """Entity-parameterize the supplied Realty baseline without rewriting its clauses."""
    text = re.sub(r"Homix Realty Inc\.?", "{{LEGAL_NAME}}", text, flags=re.IGNORECASE)
    text = re.sub(
        r"Si Zhang, CEO, on behalf of \{\{LEGAL_NAME\}\}",
        "{{BROKER_NAME}}, {{BROKER_TITLE}}, on behalf of {{LEGAL_NAME}}",
        text,
        flags=re.IGNORECASE,
    )
    return text


def add_baseline_sections(doc: Document, baseline: dict, group: str, section_numbers: list[str]) -> None:
    for number in section_numbers:
        section = baseline[group][number]
        add_heading(doc, f"{number}. {section['title']}", 2)
        for paragraph in section["paragraphs"]:
            paragraph = source_legal_text(paragraph)
            if group == "ica" and number == "II" and paragraph.startswith("4)"):
                paragraph = "[[REALTY_ONLY]]" + paragraph
            if group == "ica" and number == "III" and paragraph.startswith("C."):
                paragraph = paragraph.replace(
                    "(2) fees assessed by OneKey Multiple Listing Service for failure to comply with submission requirements (e.g., photo submission late fees, etc.),",
                    "(2) {{LISTING_SERVICE_COST_ITEM}},",
                )
            add_para(doc, paragraph)


def build_agent_master(path: Path) -> None:
    baseline = load_json_yaml(AGENT_BASELINE_PATH)
    doc = Document()
    configure_document(doc, "AGENT AFFILIATION AGREEMENT | {{LEGAL_NAME}}")

    add_title_block(doc, "AGENT AFFILIATION AGREEMENT", "{{LEGAL_NAME}} | Version {{AGENT_VERSION}} | New York")
    add_para(
        doc,
        "This agreement package combines the Commission Agreement, Commission Reporting Guideline "
        "Acknowledgement, Independent Contractor Agreement, and Non-Disclosure Agreement used in the "
        "Company's supplied enrollment package. {{LEGAL_NAME}} is the sole contracting brokerage in this release.",
    )
    add_key_values(doc, [
        ("Legal company", "{{LEGAL_NAME}}"), ("Company address", "{{ADDRESS}}"),
        ("Agent legal name", "[agent_name]"), ("Portal agent ID", "[agent_id]"),
        ("Login / contact email", "[agent_email]"), ("Phone", "[agent_phone]"),
        ("License number", "[license_number]"), ("Practice", "[practice]"),
        ("Selected plan", "[compensation_plan]"), ("Affiliation term", "[affiliation_term_months] months"),
        ("Sponsor", "[sponsor_name]"), ("Team", "[team_name]"),
    ])
    add_heading(doc, "Agreement structure")
    add_bullets(doc, [
        "Commission Agreement and plan acknowledgement.",
        "Commission Reporting Guideline Acknowledgement.",
        "Independent Contractor Agreement, Sections I-XV.",
        "Non-Disclosure Agreement, Sections I-XXIV.",
        "Realty membership application and fee disclosures only when the selected legal company requires them.",
    ])
    add_heading(doc, "Activation conditions")
    add_para(
        doc,
        "This Agreement becomes effective only after the Agent completes every required acknowledgement and "
        "signature, the Company countersigns after administrator review, required payment is recorded, and the "
        "Company approves affiliation. Portal-supplied company, plan, Sponsor, Team, term, and compensation facts "
        "are read-only. Material corrections require revalidation and, when applicable, a replacement agreement or addendum.",
    )
    add_para(
        doc,
        "Restricted tax, banking, payment-card, identity, and signing credentials are not collected in this agreement "
        "and must use an authorized restricted workflow when legally required.",
        bold=True,
    )

    new_page(doc, "Commission Agreement", "Part I | Required plan acknowledgement")
    add_para(
        doc,
        "This Commission Agreement is entered into between {{LEGAL_NAME}} (Company) and the licensed Associate "
        "identified in this agreement. When Associate performs a service on behalf of Company and Company receives "
        "compensation, Company shall pay Associate after funds are collected and have cleared into Company's account.",
    )
    add_para(doc, "Selected compensation plan (read-only): [compensation_plan]", bold=True)
    add_plan_table(doc)
    add_heading(doc, "Plan terms", 2)
    add_bullets(doc, [
        "{{PLAN_TERMS}}",
        "The selected plan applies to sales and rental transactions. There is no separate rental split plan.",
        "Sponsor Reward and Team Split are separate obligations. If the same person is both Sponsor and Team Leader, each amount is calculated and recorded separately; Sponsor Reward is calculated only from eligible Company-owned revenue under the Sponsor program.",
        "Annual or term fees are due at signing or renewal, are non-refundable, and are paid through secure Portal checkout or a Company-verified offline payment record. The selected plan remains locked for its term unless Company approves a prospective written change.",
        "If pending transactions remain when affiliation ends, the original enrollment rule applies: a 15% fee is charged on each such pending transaction, subject to applicable law and the frozen transaction facts.",
    ])
    add_acknowledgement(doc, "I understand and accept the selected plan, fee, Company Cap, Team terms, Sponsor treatment, and payment conditions.")
    add_signature_grid(doc, "AGENT SIGNATURE", "COMPANY COUNTERSIGNATURE")
    add_para(doc, "Agent anniversary / plan reset date: ________________________________", small=True)

    new_page(doc, "Commission Reporting Guideline Acknowledgement", "Part II | Required reporting acknowledgement")
    add_para(doc, source_legal_text(baseline["reporting_guideline"]))
    add_acknowledgement(doc, "I have read and agree to the Commission Reporting Guideline.")
    add_signature_grid(doc, "LICENSED SALESPERSON", "{{LEGAL_NAME}}")
    add_para(
        doc,
        "Company countersigner: {{BROKER_NAME}}, {{BROKER_TITLE}} | {{BROKER_EMAIL}}",
        small=True,
    )

    new_page(doc, "Independent Contractor Agreement", "Part III | Source-preserved Sections I-XV")
    add_key_values(doc, [
        ("Agent name", "[agent_name]"), ("Agent address", "[agent_address]"),
        ("License number / UID", "[license_number]"), ("Effective date", "[ica_effective_date]"),
        ("Brokerage name", "{{LEGAL_NAME}}"), ("Brokerage address", "{{ADDRESS}}"),
    ])
    add_baseline_sections(doc, baseline, "ica", ["I"])

    for title, numbers in [
        ("Independent Contractor Agreement | Section II", ["II"]),
        ("Independent Contractor Agreement | Sections III-IV", ["III", "IV"]),
        ("Independent Contractor Agreement | Sections V-VI", ["V", "VI"]),
        ("Independent Contractor Agreement | Section VII", ["VII"]),
        ("Independent Contractor Agreement | Sections VIII-IX", ["VIII", "IX"]),
        ("Independent Contractor Agreement | Sections X-XI", ["X", "XI"]),
        ("Independent Contractor Agreement | Sections XII-XIV", ["XII", "XIII", "XIV"]),
    ]:
        new_page(doc, title, "Part III | Original enrollment language")
        add_baseline_sections(doc, baseline, "ica", numbers)

    new_page(doc, "Independent Contractor Agreement | Section XV and Execution", "Part III | Execution")
    add_baseline_sections(doc, baseline, "ica", ["XV"])
    add_heading(doc, "Portal attribution and version record", 2)
    add_para(
        doc,
        "For each transaction, Portal preserves the effective legal company, compensation plan, plan term, "
        "participant shares, Company Dollar and cap status, Team Compensation Configuration, Team source, "
        "Team Split, Sponsor attribution, and Sponsor Reward as separate frozen facts. Corrections require "
        "an auditable Company process and do not silently rewrite an executed agreement or settled transaction.",
    )
    add_acknowledgement(doc, "I accept the complete Independent Contractor Agreement, including the Portal attribution and version record.")
    add_signature_grid(doc, "AGENT SIGNATURE", "BROKERAGE SIGNATURE")

    new_page(doc, "Non-Disclosure Agreement", "Part IV | Source-preserved Sections I-XXIV")
    add_para(
        doc,
        "This Confidentiality and Non-Disclosure Agreement (the \"Agreement\") is entered into by and between "
        "the undersigned individual (\"Recipient\") and {{LEGAL_NAME}} (\"Owner\") as of the date of signing "
        "(the \"Effective Date\").",
    )
    add_para(
        doc,
        "The purpose of this Agreement is to protect all confidential, proprietary, and sensitive information "
        "of the Owner, including but not limited to internal training content, media, customer data, "
        "communications, and technology assets.",
    )
    add_para(
        doc,
        "FOR GOOD CONSIDERATION, and in consideration of Recipient's receiving confidential information from "
        "the Owner (as defined herein), Recipient hereby agrees and acknowledges:",
    )
    add_baseline_sections(doc, baseline, "nda", ["I"])

    for title, numbers in [
        ("Non-Disclosure Agreement | Sections II-III", ["II", "III"]),
        ("Non-Disclosure Agreement | Sections IV-VIII", ["IV", "V", "VI", "VII", "VIII"]),
        ("Non-Disclosure Agreement | Sections IX-XIII", ["IX", "X", "XI", "XII", "XIII"]),
        ("Non-Disclosure Agreement | Sections XIV-XIX", ["XIV", "XV", "XVI", "XVII", "XVIII", "XIX"]),
    ]:
        new_page(doc, title, "Part IV | Original enrollment language")
        add_baseline_sections(doc, baseline, "nda", numbers)

    new_page(doc, "Non-Disclosure Agreement | Sections XX-XXIV and Execution", "Part IV | Execution")
    add_baseline_sections(doc, baseline, "nda", ["XX", "XXI", "XXII", "XXIII", "XXIV"])
    add_acknowledgement(doc, "I accept the complete Non-Disclosure Agreement.")
    add_signature_grid(doc, "RECIPIENT SIGNATURE", "AUTHORIZED COMPANY SIGNATURE")


    doc.core_properties.title = "Agent Affiliation Agreement master"
    doc.core_properties.subject = "Source-preserved two-entity Homix Agent agreement master"
    doc.core_properties.author = "Homix"
    doc.core_properties.comments = (
        "ICA Sections I-XV and NDA Sections I-XXIV derive from the supplied Realty enrollment package. "
        "Only the redlines documented in contracts/LEGAL_REVIEW_CHANGELOG.md are intentional."
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def build_team_leader_master(path: Path) -> None:
    doc = Document()
    configure_document(doc, "TEAM LEADER AGREEMENT | {{LEGAL_NAME}}")
    add_title_block(doc, "TEAM LEADER AGREEMENT", "{{LEGAL_NAME}} | Version {{TEAM_LEADER_VERSION}} | New York")
    add_para(doc, "This Team Leader Agreement supplements Team Leader's completed Agent Affiliation Agreement for {{LEGAL_NAME}} and establishes an internal leadership role, not a separate brokerage, employment relationship, partnership, franchise, ownership interest, or authority to bind Company. The Team and every licensed Team Member must be associated with the same licensed brokerage.")
    add_key_values(doc, [
        ("Legal company", "{{LEGAL_NAME}}"), ("Company address", "{{ADDRESS}}"),
        ("Team Leader", "[agent_name]"), ("License number", "[license_number]"),
        ("Team name", "[team_name]"), ("Expected members", "[expected_member_count]"),
        ("Team positioning", "[team_positioning]"), ("Required plan", "[compensation_plan]"),
    ])
    add_legal_section(doc, "I. Eligibility, appointment, and activation", [
        "Appointment requires an Active Agent account, completed onboarding with the same Company, Solo Pro, an approved Team Leader application, completion of this Agreement, and Company countersignature. Company may review experience, license standing, proposed recruiting, Team positioning, expected members, supervision capacity, payment standing, compliance history, and business readiness.",
        "Portal records the Team in forming status until all activation conditions are satisfied. After Company countersigns, Portal may permit entity-locked recruiting links. The Team becomes Active only after at least one approved Team Member completes the applicable Agent Agreement and Company verifies every activation condition. Company may pause, deny, or revoke activation for compliance, licensing, payment, inactivity, risk, or operational reasons.",
        "Team Leader's authority is limited to the internal role described here. Only Broker and authorized Company administrators may approve affiliation, sponsorship, legal status, advertising, commission policy, exceptions, final transaction compliance, or Company obligations.",
    ])
    add_legal_section(doc, "II. Legal Team boundary", [
        "Team Leader.companyId, Team.companyId, and every Team Member.companyId must match. The Team may not recruit, advertise, or hold out a licensee associated with another brokerage as a legal Team Member. A recruiting link locks the licensed company and does not permit the candidate to select another company.",
        "Shared branding across Homix licensed companies may use a non-legal internal team_group only when each legal Team, agreement, compensation configuration, member roster, and transaction remains separate under the correct licensed company. A shared Broker does not merge the companies or permit cross-company supervision or compensation records to be treated as one legal Team.",
    ])

    new_page(doc, "Team Compensation Configuration", "Part II | Required acknowledgement")
    add_para(doc, "Team Leader personal-production plan (read-only): [compensation_plan]", bold=True)
    add_key_values(doc, [
        ("Team name", "[team_name]"), ("Terms effective", "[team_terms_effective_from]"),
        ("Standard Team Split", "[team_split_pct]% of Agent Side"),
        ("Team-sourced Split", "[team_sourced_split_pct]% of Agent Side"),
        ("Annual Team Cap", "[team_cap_usd]"), ("Configuration version", "[team_config_version]"),
    ])
    add_legal_section(doc, "Calculation order", [
        "For Team Member transactions, approved outside referral or source economics are applied first, Company Dollar next, and Team Split from Agent Side after Company Dollar. Company Cap and Team Cap are separate ledgers and may be reached at different times. Agent shares and source facts are applied before a Team distribution.",
        "Sponsor Reward and Team Split are separate financial obligations. If Team Leader is also the qualifying Sponsor, Portal records and reports both independently. Sponsor Reward is calculated from eligible Company-owned revenue under the current Sponsor program and does not reduce Agent commission or Company Cap credit. Team Split is calculated under the frozen Team Compensation Configuration.",
        "No Team Split, Sponsor Reward, cap credit, lead, commission, Team growth, or continued program availability is guaranteed. Payments remain subject to closed transactions, cleared funds, compliance approval, frozen facts, offsets, and applicable policy.",
    ])
    add_legal_section(doc, "Versioning and change control", [
        "The configuration is effective-dated and prospective. An accepted, executed, effective, or transaction-used version may not be edited in place. A change requires a new Portal version, effective date, audit history, policy validation, Company approval, and any acceptance or addendum required for affected members.",
        "Team Leader may propose Team Split and Team Cap values within the ranges made available in Portal. Company retains final approval because these terms affect Company financial administration, Agent agreements, compliance, and future payment obligations. Company may reject, modify before publication, pause, or prospectively replace a proposal. Company will not silently rewrite an accepted version.",
    ])
    add_acknowledgement(doc, "I accept the locked Team configuration, understand that my personal plan is Solo Pro, and understand that Sponsor Reward and Team Split are independent obligations.")
    add_para(doc, "Team Leader initials: __________", bold=True)

    new_page(doc, "Recruiting, Sponsor Attribution, and Onboarding", "Part III | Team growth")
    add_legal_section(doc, "I. Recruiting links and representations", [
        "A Team recruiting link locks Company, Team, Team Member plan, current compensation configuration, and Sponsor. Sponsor defaults to Team Leader but may be another Active Agent on the same Team who actually recruited the candidate. Team Leader may not select an outside person, another Team's Agent, or an inactive Agent as Sponsor.",
        "Team Leader will not promise affiliation, license sponsorship, approval, guaranteed income, guaranteed leads, a closing, a cap result, unapproved compensation, a private referral payment, or a Sponsor reward outside the Company program. Recruiting materials must accurately identify {{LEGAL_NAME}}, the Team, plan, fees, requirements, and the conditional nature of approval.",
        "A personally referred candidate may request this Team while retaining the original Sponsor. Team Leader's acceptance of the Team request does not replace Sponsor attribution and does not substitute for Company approval. If Team Leader declines the request, the candidate returns to plan selection and Sponsor remains unchanged.",
    ])
    add_legal_section(doc, "II. Candidate review and Team consent", [
        "Team Leader may review only the candidate information and onboarding status reasonably necessary to decide a Team request and support onboarding. Team Leader may accept or decline membership but may not activate the Company account, countersign Company agreements, waive payment, change the legal company, edit license records, or approve a compliance exception.",
        "An accepted Team request freezes the current Team Compensation Configuration for candidate acceptance. Candidate must separately accept Team terms and sign the same-entity Agent Agreement. Material corrections before eSign require an auditable event showing the old value, new value, operator, reason, and authorization. After eSign preparation, company, plan, Team, Sponsor, term, or compensation changes require cancellation and a new agreement or authorized addendum.",
    ])
    add_legal_section(doc, "III. Onboarding support", [
        "Team Leader will provide commercially reasonable orientation, systems guidance, training direction, role expectations, transaction-file instruction, and escalation support. Team Leader will monitor progress but will not request SSN, full bank data, card information, W-9 contents, identity credentials, or other restricted information outside approved workflows.",
        "Team activation and membership remain subject to completed agreement signatures, required payments or verified offline-payment evidence, license and company checks, required association steps, Company review, and final approval. Company may deny or condition approval notwithstanding Team Leader acceptance.",
    ])

    new_page(doc, "Leadership, Supervision, and Compliance", "Part IV | Operating duties")
    add_legal_section(doc, "I. Leadership and member support", [
        "Team Leader will provide commercially reasonable coaching, workflow supervision, transaction escalation, training, accountability, and operational support consistent with Company policy. Team Leader will maintain accurate member status, avoid unauthorized compensation promises, and promptly tell Company when a member leaves, becomes inactive, changes company, or presents a material risk.",
        "Team Leader will not create an employment relationship with a Team Member on Company's behalf, hold client funds, issue Company checks, sign Company contracts, modify Company records outside granted permissions, or represent that the Team is an independent brokerage or separate legal entity.",
    ])
    add_legal_section(doc, "II. Brokerage supervision and escalation", [
        "Broker retains final authority over licensing sponsorship, legal supervision, compliance, advertising, transactions, Company compensation policy, Team configuration ranges, Team status, access, discipline, and client-protection decisions. Team Leader's day-to-day coaching assists but never replaces Broker supervision.",
        "Team Leader will promptly escalate license issues, complaints, claims, fair housing concerns, agency conflicts, escrow or funds issues, advertising violations, incomplete or inaccurate files, suspected fraud, privacy or cybersecurity incidents, threats, discrimination, unauthorized practice of law, and any matter requiring brokerage oversight.",
        "Team Leader will not give legal or tax advice unless independently qualified and expressly authorized. Team Leader will direct legal questions to counsel and compliance questions to Broker or an authorized Company administrator.",
    ])
    add_legal_section(doc, "III. Advertising and Team identity", [
        "Team names, websites, domains, social accounts, listing marketing, signs, and advertisements require accurate brokerage identification and Company approval where required. Team Leader will ensure compliance with Article 12-A, 19 NYCRR Part 175, fair housing, intellectual-property law, Company brand standards, and applicable professional rules.",
        "The HOMIX name and approved Team branding are licensed only during the authorized role. On suspension, termination, dissolution, or Company request, Team Leader will stop affected use, preserve records, and transfer or relinquish Company-controlled domains, handles, files, and access as directed.",
    ])

    new_page(doc, "Data, Records, and Risk Allocation", "Part V | Security and accountability")
    add_legal_section(doc, "I. Data access and confidentiality", [
        "Team Leader access is limited to information reasonably needed to manage the Team. Team Leader may not access W-9 forms, ACH or bank details, payment-card data, administrator notes, sensitive evidence, Social Security numbers, restricted identity information, or unrelated Agent records except through an authorized role, documented purpose, and approved workflow.",
        "Team Leader will protect candidate, member, client, transaction, compensation, training, and Company information under the Agent NDA and Company security policies. Team Leader will not export a Team roster, share credentials, place protected data into unapproved systems, or use Team information for an outside business. Duties continue after the role ends.",
    ])
    add_legal_section(doc, "II. Records and audit cooperation", [
        "Team Leader will keep recruiting, Team consent, coaching, compensation proposal, advertising, complaint, and escalation records in approved Company systems. Team Leader will cooperate with Company audits, investigations, payment reconciliation, transaction review, legal holds, and regulator requests and will not delete, alter, or conceal required records.",
        "Sponsor attribution, Team Split, Company Dollar, Sponsor Reward, Company Cap, Team Cap, and Agent net must remain distinct ledger entries. Team Leader will promptly report a suspected duplicate, wrong Sponsor, wrong Team, wrong plan, or calculation error and will not resolve it through an off-platform payment or undocumented side agreement.",
    ])
    add_legal_section(doc, "III. Insurance, indemnity, and expenses", [
        "The insurance, indemnification, legal-cost, tax, and independent-contractor provisions of Team Leader's Agent Affiliation Agreement apply to Team Leader activity. Team Leader is responsible for expenses voluntarily incurred for the Team unless Company approves reimbursement in writing and may not incur an obligation in Company's name.",
        "To the extent permitted by law, Team Leader is responsible for attributable loss caused by intentional misconduct, fraud, knowing legal violation, unauthorized promise, misuse of funds or data, or material breach of this Agreement. Nothing shifts Company's non-delegable brokerage duties or liability that cannot lawfully be shifted.",
    ])

    new_page(doc, "Term, Team Changes, and Dissolution", "Part VI | Continuity")
    add_legal_section(doc, "I. Term and role changes", [
        "This Agreement begins only after all required signatures, Company approval, and Portal verification. It continues until Team Leader resigns, Company removes Team Leader, the Team dissolves, the underlying Agent affiliation ends, or another event stated in policy occurs. Ending Team Leader status does not by itself terminate Agent affiliation, but ending Agent affiliation immediately ends Team Leader authority.",
        "Membership changes, transfers, Team Leader replacement, Team merger, Team closure, legal-company change, and material compensation changes are prospective and require Portal records, Company approval, and any agreement or addendum required for affected people. Historical transactions remain governed by frozen facts and versions.",
    ])
    add_legal_section(doc, "II. Suspension and termination", [
        "Company may suspend recruiting, access, proposed compensation changes, or Team operations while reviewing compliance, licensing, payment, inactivity, security, client-protection, or operational concerns. Company may require corrective action, training, records, or supervision before restoring privileges.",
        "On resignation, removal, or termination, Team Leader will stop representing the leadership role, return Company property and access, preserve records, cooperate in member and client transition, and avoid interfering with Company supervision. Pending compensation and Sponsor rewards remain subject to frozen transaction facts, this Agreement, the Agent Agreement, and lawful offsets.",
    ])
    add_legal_section(doc, "III. Team dissolution and successor", [
        "A Team is not owned separately from the brokerage relationship and cannot continue as a legal Company Team without an approved Active Team Leader. Company may appoint an interim or successor leader, move members to Solo or another accepted Team prospectively, or dissolve the Team after notice and required consents. Client choice and legal obligations control all client and transaction transitions.",
        "Team branding, domains, social accounts, shared files, lead records, templates, and Company-provided systems will be handled under Company ownership, licenses, written policy, and applicable law. No dissolution changes an Agent's separate lawful Agent Data rights or guarantees transfer of a client agreement.",
    ])
    add_legal_section(doc, "IV. Survival and general terms", [
        "Confidentiality, data security, record retention, payment reconciliation, intellectual property, indemnity, transition, and provisions that by their nature survive remain effective. This Agreement, Agent Affiliation Agreement, frozen Team Compensation Configuration, Team Member agreements, written policies, and signed addenda govern the role. New York law governs, severability applies, and no oral statement modifies these terms.",
    ])

    new_page(doc, "Team Leader Agreement - Execution", "Part VII | Required signatures")
    add_heading(doc, "Final acknowledgement", 2)
    add_para(doc, "Team Leader confirms that the legal Company, required Solo Pro plan, Team configuration, Sponsor rules, activation conditions, role boundaries, data restrictions, and Company oversight were disclosed. Company countersigns manually only after administrator review. Recruiting privileges are not effective until Portal verifies all signatures and Company approval.")
    add_acknowledgement(doc, "I reviewed and accept the complete Team Leader Agreement and the locked Team Compensation Configuration.")
    add_key_values(doc, [
        ("Contracting company", "{{LEGAL_NAME}}"), ("Company address", "{{ADDRESS}}"),
        ("Team Leader", "[agent_name]"), ("License", "[license_number]"),
        ("Team", "[team_name]"), ("Configuration", "[team_config_version]"),
    ])
    add_signature_grid(doc, "TEAM LEADER SIGNATURE", "COMPANY COUNTERSIGNATURE")
    add_para(doc, "Countersigner: {{BROKER_NAME}}, {{BROKER_TITLE}} | {{BROKER_EMAIL}} | {{LEGAL_NAME}}", small=True)
    doc.core_properties.title = "Team Leader Agreement master"
    doc.core_properties.subject = "Two-entity Homix Team Leader agreement master"
    doc.core_properties.author = "Homix"
    doc.core_properties.comments = "Team Leader must be Solo Pro and associated with the Team's licensed company."
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def apply_realty_paragraphs(doc: Document, include: bool) -> None:
    for paragraph in list(doc.paragraphs):
        if "[[REALTY_ONLY]]" not in paragraph.text:
            continue
        if include:
            for run in paragraph.runs:
                run.text = run.text.replace("[[REALTY_ONLY]]", "")
        else:
            paragraph._element.getparent().remove(paragraph._element)


def iter_paragraphs(parent):
    yield from parent.paragraphs
    for table in parent.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def replace_placeholders(doc: Document, replacements: dict[str, str]) -> None:
    containers = [doc]
    for section in doc.sections:
        containers.extend([section.header, section.footer])
    for container in containers:
        for paragraph in iter_paragraphs(container):
            for run in paragraph.runs:
                original = run.text
                updated = original
                for key, value in replacements.items():
                    updated = updated.replace(f"{{{{{key}}}}}", value)
                if updated != original:
                    run.text = updated


def clear_esign_placeholders(doc: Document) -> None:
    """Keep merge keys in the masters, but leave publishable PDF fields blank."""
    pattern = re.compile(
        r"\[[a-z][a-z0-9_]*\](?:% of Agent Side| months)?",
        flags=re.IGNORECASE,
    )
    for paragraph in iter_paragraphs(doc):
        for run in paragraph.runs:
            original = run.text
            updated = pattern.sub("", original)
            if updated != original:
                run.text = updated


def generate_entity_docx(
    master: Path,
    destination: Path,
    entity: dict,
    is_agent: bool,
    plan_key: str | None = None,
) -> None:
    doc = Document(master)
    if is_agent:
        apply_realty_paragraphs(doc, bool(entity["requires_libor_onekey"]))
    replacements = {
        "LEGAL_NAME": entity["legal_name"], "ADDRESS": entity["address"],
        "BROKER_NAME": entity["broker_name"], "BROKER_TITLE": entity["broker_title"],
        "BROKER_EMAIL": entity["broker_email"], "AGENT_VERSION": entity["agent_version"],
        "TEAM_LEADER_VERSION": entity["team_leader_version"],
        "LISTING_SERVICE_COST_ITEM": (
            "fees assessed by OneKey Multiple Listing Service for failure to comply with submission "
            "requirements (e.g., photo submission late fees, etc.)"
            if entity["requires_libor_onekey"]
            else "third-party listing-service or association fees applicable to Agent's practice"
        ),
    }
    if is_agent:
        if plan_key not in AGENT_PLANS:
            raise ValueError(f"Unknown Agent plan: {plan_key}")
        plan = AGENT_PLANS[plan_key]
        replacements.update({
            "PLAN_LABEL": plan["label"],
            "PLAN_ECONOMICS": plan["economics"],
            "PLAN_FEE": plan["fee"],
            "PLAN_CAP_RULE": plan["cap_rule"],
            "PLAN_TERMS": plan["terms"],
        })
    replace_placeholders(doc, replacements)
    clear_esign_placeholders(doc)
    plan_title = f" {AGENT_PLANS[plan_key]['label']}" if is_agent and plan_key else ""
    doc.core_properties.title = (
        f"{entity['legal_name']}{plan_title} "
        f"{'Agent Affiliation' if is_agent else 'Team Leader'} Agreement"
    )
    doc.core_properties.author = entity["legal_name"]
    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destination)


def convert_to_pdf(docx_path: Path) -> Path:
    office = shutil.which("soffice") or shutil.which("libreoffice")
    if not office:
        raise RuntimeError("LibreOffice is required to generate release-candidate PDFs.")
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    path = PDF_DIR / f"{docx_path.stem}.pdf"
    path.unlink(missing_ok=True)
    with tempfile.TemporaryDirectory(prefix="homix-libreoffice-") as profile:
        result = subprocess.run(
            [
                office,
                f"-env:UserInstallation={Path(profile).as_uri()}",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(PDF_DIR),
                str(docx_path),
            ],
            capture_output=True,
            text=True,
            check=False,
        )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice conversion failed: {result.stderr or result.stdout}")
    if not path.exists():
        raise RuntimeError(f"LibreOffice did not create {path}")
    return path


def insert_realty_appendices(path: Path) -> None:
    if not LIBOR_APPLICATION_PATH.exists():
        raise RuntimeError(f"Missing official LIBOR application attachment: {LIBOR_APPLICATION_PATH}")
    if not REALTY_FEE_DISCLOSURE_PATH.exists():
        raise RuntimeError(f"Missing Realty fee disclosure attachment: {REALTY_FEE_DISCLOSURE_PATH}")
    source = PdfReader(str(path))
    application = PdfReader(str(LIBOR_APPLICATION_PATH))
    disclosures = PdfReader(str(REALTY_FEE_DISCLOSURE_PATH))
    if len(application.pages) != 1:
        raise RuntimeError("The official LIBOR application attachment must contain exactly one page.")
    if len(disclosures.pages) != 2:
        raise RuntimeError("The Realty fee disclosure attachment must contain exactly two pages.")
    writer = PdfWriter()
    for page in source.pages:
        writer.add_page(page)
    writer.add_page(application.pages[0])
    for page in disclosures.pages:
        writer.add_page(page)
    temp_path = path.with_suffix(".merged.pdf")
    with temp_path.open("wb") as handle:
        writer.write(handle)
    temp_path.replace(path)


def canonicalize_static_pdf(path: Path, title: str, author: str) -> None:
    """Remove converter timestamps so unchanged source produces an identical PDF hash."""
    source = PdfReader(str(path))
    writer = PdfWriter()
    for page in source.pages:
        writer.add_page(page)
    writer.add_metadata({
        "/Title": title,
        "/Author": author,
        "/Creator": "Homix contract generator",
        "/Producer": "pypdf",
    })
    temp_path = path.with_suffix(".canonical.pdf")
    with temp_path.open("wb") as handle:
        writer.write(handle)
    temp_path.replace(path)


def normalized_contract_words(text: str) -> list[str]:
    return re.findall(
        r"[a-z0-9]+",
        text.lower().replace("’", "'").replace("“", '"').replace("”", '"'),
    )


def assert_source_legal_fidelity(text: str, entity: dict) -> None:
    baseline = load_json_yaml(AGENT_BASELINE_PATH)
    actual = " ".join(normalized_contract_words(text))
    missing: list[str] = []
    checked = 0
    for group in ("ica", "nda"):
        for number, section in baseline[group].items():
            for paragraph in section["paragraphs"]:
                if (
                    group == "ica"
                    and number == "II"
                    and paragraph.startswith("4)")
                    and not entity["requires_libor_onekey"]
                ):
                    continue
                expected = source_legal_text(paragraph)
                expected = expected.replace("{{LEGAL_NAME}}", entity["legal_name"])
                expected = expected.replace("{{BROKER_NAME}}", entity["broker_name"])
                expected = expected.replace("{{BROKER_TITLE}}", entity["broker_title"])
                if group == "ica" and number == "III" and paragraph.startswith("C."):
                    expected = expected.replace(
                        "(2) fees assessed by OneKey Multiple Listing Service for failure to comply with "
                        "submission requirements (e.g., photo submission late fees, etc.),",
                        "(2) "
                        + (
                            "fees assessed by OneKey Multiple Listing Service for failure to comply with "
                            "submission requirements (e.g., photo submission late fees, etc.)"
                            if entity["requires_libor_onekey"]
                            else "third-party listing-service or association fees applicable to Agent's practice"
                        )
                        + ",",
                    )
                words = normalized_contract_words(expected)
                if not words:
                    continue
                checked += 1
                sample_size = min(12, len(words))
                head = " ".join(words[:sample_size])
                tail = " ".join(words[-sample_size:])
                if head not in actual or tail not in actual:
                    missing.append(f"{group.upper()} {number}: {head}")
    if missing:
        preview = "; ".join(missing[:8])
        raise RuntimeError(
            f"Source legal fidelity failed for {entity['legal_name']} "
            f"({len(missing)}/{checked} clauses): {preview}"
        )


def assert_clean_pdf(
    path: Path,
    entity: dict,
    is_agent: bool,
    plan_key: str | None = None,
) -> dict:
    reader = PdfReader(str(path))
    expected = 21 if is_agent and entity["requires_libor_onekey"] else 18 if is_agent else 7
    if len(reader.pages) != expected:
        raise RuntimeError(f"{path.name} has {len(reader.pages)} pages; expected {expected}.")
    if reader.is_encrypted or reader.get_fields():
        raise RuntimeError(f"Encrypted or interactive output: {path}")
    root = reader.trailer["/Root"]
    if root.get("/JavaScript") or root.get("/OpenAction"):
        raise RuntimeError(f"Unexpected document action: {path}")
    for index, page in enumerate(reader.pages, 1):
        if page.get("/Annots"):
            raise RuntimeError(f"Unexpected annotations on {path.name}, page {index}")
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    normalized_text = re.sub(r"\s+", " ", text).strip().lower()
    if re.search(r"\[[a-z][a-z0-9_]*\]", normalized_text):
        raise RuntimeError(f"{path.name} contains a printed eSign placeholder.")
    other = "Homix Living Inc." if entity["legal_name"] == "Homix Realty Inc." else "Homix Realty Inc."
    if other.lower() in normalized_text:
        raise RuntimeError(f"{path.name} contains the other contracting entity.")
    if entity["legal_name"].lower() not in normalized_text or entity["address"].lower() not in normalized_text:
        raise RuntimeError(f"{path.name} is missing the selected entity or address.")
    if not entity["requires_libor_onekey"] and any(term in normalized_text for term in ("libor", "onekey", "mls")):
        raise RuntimeError(f"Living release {path.name} contains Realty-only language.")
    if is_agent:
        if plan_key not in AGENT_PLANS:
            raise RuntimeError(f"{path.name} has no frozen Agent plan.")
        plan = AGENT_PLANS[plan_key]
        if plan["terms"].lower() not in normalized_text:
            raise RuntimeError(f"{path.name} is missing the frozen {plan['label']} terms.")
        assert_source_legal_fidelity(text, entity)
        if any(term in normalized_text for term in ("legal-review candidate", "credit card authorization", "card number:", "cvv:")):
            raise RuntimeError(f"{path.name} contains internal-review or prohibited payment-card content.")
        if entity["requires_libor_onekey"]:
            official = PdfReader(str(LIBOR_APPLICATION_PATH)).pages[0].get_contents().get_data()
            inserted = reader.pages[18].get_contents().get_data()
            if hashlib.sha256(inserted).digest() != hashlib.sha256(official).digest():
                raise RuntimeError(f"{path.name} is missing the official LIBOR Rev 10/25 application page.")
            disclosures = PdfReader(str(REALTY_FEE_DISCLOSURE_PATH))
            for offset, disclosure in enumerate(disclosures.pages, start=19):
                expected_content = disclosure.get_contents().get_data()
                actual_content = reader.pages[offset].get_contents().get_data()
                if hashlib.sha256(actual_content).digest() != hashlib.sha256(expected_content).digest():
                    raise RuntimeError(f"{path.name} is missing Realty disclosure page {offset - 18}.")
    return {"file": path.name, "pages": len(reader.pages),
            "sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
            "entity": entity["legal_name"], "agreement": "agent" if is_agent else "team_leader",
            "plan": plan_key if is_agent else None}


def write_release_index(records: list[dict]) -> None:
    (PDF_DIR / "release-index.json").write_text(json.dumps({"contracts": records}, indent=2) + "\n", encoding="utf-8")
    lines = ["# Homix onboarding contract release candidates", "",
             "Generated from two legal-reviewable DOCX masters and `contracts/entities.yml`.",
             "Homix approved this release set. The `-candidate` suffix is retained only for release naming compatibility; the files below are the approved sources for immutable production eSign releases.", "",
             "| File | Entity | Agreement | Plan | Pages | SHA-256 |", "| --- | --- | --- | --- | ---: | --- |"]
    for record in records:
        lines.append(
            f"| `{record['file']}` | {record['entity']} | {record['agreement']} | "
            f"{record.get('plan') or '—'} | {record['pages']} | `{record['sha256']}` |"
        )
    lines += ["", "Do not overwrite a published eSign version. Any PDF or field-manifest change requires a new immutable template version, schema hash, Portal pin, and smoke test.", ""]
    (PDF_DIR / "README.md").write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--author-masters", action="store_true")
    args = parser.parse_args()
    if args.author_masters:
        build_agent_master(AGENT_MASTER)
        build_team_leader_master(TEAM_LEADER_MASTER)
    if not AGENT_MASTER.exists() or not TEAM_LEADER_MASTER.exists():
        raise SystemExit("Missing DOCX masters. Run once with --author-masters.")
    if not MANIFEST_PATH.exists():
        raise SystemExit(f"Missing field manifest: {MANIFEST_PATH}")
    records = []
    for entity_key, entity in load_json_yaml(ENTITIES_PATH).items():
        entity_filename = "Homix_Realty" if entity_key == "homix_realty" else "Homix_Living"
        leader_docx = GENERATED_DIR / f"{entity['team_leader_filename']}.docx"
        generate_entity_docx(TEAM_LEADER_MASTER, leader_docx, entity, False)
        for plan_key, plan in AGENT_PLANS.items():
            agent_filename = (
                f"{entity_filename}_{plan['filename']}_Agent_Affiliation_Agreement_"
                f"v{entity['agent_version']}"
            )
            agent_docx = GENERATED_DIR / f"{agent_filename}.docx"
            generate_entity_docx(AGENT_MASTER, agent_docx, entity, True, plan_key)
            agent_pdf = convert_to_pdf(agent_docx)
            if entity["requires_libor_onekey"]:
                insert_realty_appendices(agent_pdf)
            canonicalize_static_pdf(
                agent_pdf,
                f"{entity['legal_name']} {plan['label']} Agent Affiliation Agreement "
                f"{entity['agent_version']}",
                entity["legal_name"],
            )
            records.append(assert_clean_pdf(agent_pdf, entity, True, plan_key))
        leader_pdf = convert_to_pdf(leader_docx)
        canonicalize_static_pdf(
            leader_pdf,
            f"{entity['legal_name']} Team Leader Agreement {entity['team_leader_version']}",
            entity["legal_name"],
        )
        records.append(assert_clean_pdf(leader_pdf, entity, False))
    write_release_index(records)
    for record in records:
        print(f"{record['file']}: {record['pages']} pages {record['sha256']}")


if __name__ == "__main__":
    main()
